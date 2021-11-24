""" Модуль 
"""

import docx
import argparse
import datetime
from docx.shared import Pt,Cm 
 
def paragraph_replace(paragraph,old,new,style):
    """ Функция ищет образец  в параграфе и заменят его новым значением
    """

    if old in paragraph.text:
        old1 = '1'+old
        old2 = '2'+old
        old3 = '3'+old
        if (not old1 in paragraph.text) and \
                (not old2 in paragraph.text) and \
                (not old3 in paragraph.text):
            paragraph.text = paragraph.text.replace(old,new)
            paragraph.style = style
                
                
    return paragraph

def doc_replace(document,old,new,style):
    """ Функция ищет образец в документе и заменяет его новым
    """

    for tb in document.tables: 
        for column in tb.columns:
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    paragraph = paragraph_replace(paragraph,old,new,style)
    return document

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("year", type=int,
                    help="make planning for YEAR")
    
    args = parser.parse_args()
    b = datetime.date(2018,12,31)
    numdays = 2*369
    # numdays = 7
    # Читаем образцы четные и не четные листы
    document = docx.Document("templ1.docx")
    document1 = docx.Document("templ2.docx")
    # Меням два стиля для этих двух документов
    style1 = document1.styles['Normal']
    font1 = style1.font
    font1.name = 'new times roman'
    font1.size = Pt(9)

    style = document.styles['Normal']
    font = style.font
    font.name = 'new times roman'
    font.size = Pt(9)


    month_list_o = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь',
           'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    month_list = ['Января', 'Февраля', 'Марта', 'Апреля', 'Мая', 'Июня',
           'Июля', 'Августа', 'Сентября', 'Октября', 'Ноября', 'Декабря']
    # Задаем дату первого дня года, для которого делается планнинг
    n = datetime.date(args.year,1,1)
    # определяем дату понедельника в недель, которая содержит 1 января
    delta = datetime.timedelta(days=n.weekday())
    monday = n-delta
    # Пробегая по всем дням года меняем в образцах старые значения
    # дней на новые
    for i in range(0,numdays):
        delta = datetime.timedelta(days=i)
        old_day = b+delta
        new_day = monday+delta
        s = str(old_day.day)+' '+month_list_o[old_day.month-1]+\
            ' '+str(old_day.year)
        s1 = str(new_day.strftime('%d'))+' '+\
            month_list[new_day.month-1]+' '+str(new_day.year)
        document = doc_replace(document,s,s1,style)
        document1 = doc_replace(document1,s,s1,style1)
        print(s1)
       
    margin_top = Cm(0.8)
    margin_botton = Cm(0.2)
    sections = document.sections
    for section in sections:
        section.top_margin = margin_top
        section.bottom_margin = margin_botton
    document.save("example.docx")


    margin_top = Cm(0.4)
    margin_botton = Cm(0.6)
    sections = document1.sections
    for section in sections:
        section.top_margin = margin_top
        section.bottom_margin = margin_botton

    document1.save("example1.docx")



